import logging
import re
import uuid
from io import BytesIO
from pathlib import Path
from typing import Any

from fastapi import APIRouter, Depends, HTTPException, Request
from sqlalchemy import desc
from sqlalchemy.orm import Session
from starlette.responses import StreamingResponse
from starlette.datastructures import UploadFile

from database.models import ChatAttachment, ChatJob, ChatMessage, ChatRecord, ChatSession, UserAccount, now_utc
from database.session import get_db
from models.schemas import ChatExportRequest, ChatJobOut, ChatModelOut, ChatRequest, ChatResponse, ChatSessionDetail, ChatSessionOut
from services.auth_service import current_user
from services.chat_context_service import load_recent_chat_history
from services.chat_model_service import list_active_chat_models, resolve_chat_model
from services.document_extract import DEFAULT_MAX_CHARS, extract_document_text
from services.openai_service import OpenAIService, OpenAIServiceError
from services.rate_limit import InMemoryRateLimiter
from services.settings_service import normalize_provider
from services.token_usage_service import record_token_usage

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api", tags=["chat"])
rate_limiter = InMemoryRateLimiter()

UPLOAD_DIR = Path(__file__).resolve().parents[1] / "uploads" / "chat"
MAX_FILES = 5
MAX_FILE_SIZE = 20 * 1024 * 1024
MAX_EXTRACTED_TEXT = DEFAULT_MAX_CHARS
ALLOWED_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".pdf",
    ".doc",
    ".docx",
    ".ppt",
    ".pptx",
    ".xls",
    ".xlsx",
    ".py",
    ".js",
    ".jsx",
    ".ts",
    ".tsx",
    ".html",
    ".css",
    ".xml",
    ".yaml",
    ".yml",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".gif",
}
IMAGE_MIME_BY_EXT = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".gif": "image/gif",
}
REQUIRE_EXTRACTED_TEXT_EXTENSIONS = {".docx", ".pdf", ".pptx", ".xlsx"}
WORD_FONT_NAME = "微软雅黑"


def make_session_title(message: str) -> str:
    title = " ".join(message.strip().split())
    return title[:42] or "New chat"


def safe_filename(filename: str) -> str:
    name = Path(filename or "attachment").name
    return re.sub(r"[^a-zA-Z0-9._-]+", "_", name)[:180] or "attachment"


def extract_text(filename: str, content_type: str, data: bytes) -> str | None:
    """Extract readable text from plain text and office documents."""
    text = extract_document_text(filename, content_type, data, max_chars=MAX_EXTRACTED_TEXT)
    if text:
        logger.info("Extracted %s chars from %s", len(text), filename)
    else:
        logger.warning("No text extracted from %s (%s)", filename, content_type)
    return text


def strip_assistant_markup(content: str) -> str:
    answer = re.search(r"<ai_answer>\s*([\s\S]*?)\s*</ai_answer>", content, flags=re.IGNORECASE)
    if answer:
        return answer.group(1).strip()
    return re.sub(r"<ai_thought_summary>[\s\S]*?</ai_thought_summary>", "", content, flags=re.IGNORECASE).strip()


def normalize_word_math_delimiters(content: str) -> str:
    return (
        content.replace("\r\n", "\n")
        .replace("\\[", "$$")
        .replace("\\]", "$$")
        .replace("\\(", "$")
        .replace("\\)", "$")
    )


def clean_markdown_text(value: str) -> str:
    text = value.strip()
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    text = text.replace("\\|", "|")
    return text


def set_run_font(run: Any, size: float | None = 10.5, bold: bool | None = None, italic: bool | None = None, color: str | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = WORD_FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph: Any, before: int = 0, after: int = 6, line_spacing: float = 1.25) -> None:
    from docx.shared import Pt

    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_inline_markdown_runs(paragraph: Any, text: str, size: float = 10.5) -> None:
    token_pattern = re.compile(r"(`[^`]+`|\*\*.+?\*\*|__.+?__|\*[^*\n]+\*|_[^_\n]+_)", re.DOTALL)
    position = 0
    text = clean_markdown_text(text)

    for match in token_pattern.finditer(text):
        if match.start() > position:
            set_run_font(paragraph.add_run(text[position : match.start()]), size=size)

        token = match.group(0)
        if token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color="374151")
        elif (token.startswith("**") and token.endswith("**")) or (token.startswith("__") and token.endswith("__")):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True)
        elif (token.startswith("*") and token.endswith("*")) or (token.startswith("_") and token.endswith("_")):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, italic=True)
        position = match.end()

    if position < len(text):
        set_run_font(paragraph.add_run(text[position:]), size=size)


def table_cells(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    cells = re.split(r"(?<!\\)\|", stripped)
    return [clean_markdown_text(cell) for cell in cells]


def is_table_separator(line: str) -> bool:
    cells = table_cells(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def is_table_start(lines: list[str], index: int) -> bool:
    return index + 1 < len(lines) and "|" in lines[index] and is_table_separator(lines[index + 1])


def shade_cell(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell: Any, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, after=0, line_spacing=1.15)
    add_inline_markdown_runs(paragraph, text, size=9.5)
    for run in paragraph.runs:
        run.bold = bold or run.bold


def add_markdown_table(document: Any, table_lines: list[str]) -> None:
    if len(table_lines) < 2:
        return

    rows = [table_cells(line) for line in table_lines[:1] + table_lines[2:]]
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.autofit = True

    for row_index, row_values in enumerate(rows):
        row = table.rows[row_index]
        for column_index in range(column_count):
            cell = row.cells[column_index]
            value = row_values[column_index] if column_index < len(row_values) else ""
            set_cell_text(cell, value, bold=row_index == 0)
            if row_index == 0:
                shade_cell(cell, "EEF2FF")
            elif row_index % 2 == 0:
                shade_cell(cell, "F8FAFC")

    document.add_paragraph()


def configure_word_document(document: Any) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    for style in document.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        if not hasattr(style, "font"):
            continue
        style.font.name = WORD_FONT_NAME
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), WORD_FONT_NAME)

    normal_style = document.styles["Normal"]
    normal_style.font.name = WORD_FONT_NAME
    normal_style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), WORD_FONT_NAME)
    normal_style.font.size = Pt(10.5)


def add_markdown_to_document(document: Any, content: str) -> None:
    lines = normalize_word_math_delimiters(content).split("\n")
    in_code_block = False
    code_lines: list[str] = []
    index = 0

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        paragraph = document.add_paragraph()
        set_paragraph_spacing(paragraph, before=4, after=8, line_spacing=1.1)
        run = paragraph.add_run("\n".join(code_lines))
        set_run_font(run, size=9.5, color="111827")
        code_lines = []

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                flush_code()
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            index += 1
            continue

        if in_code_block:
            code_lines.append(raw_line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if is_table_start(lines, index):
            table_lines = [lines[index], lines[index + 1]]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            paragraph = document.add_heading("", level=len(heading.group(1)))
            set_paragraph_spacing(paragraph, before=12, after=6, line_spacing=1.18)
            run = paragraph.add_run(clean_markdown_text(heading.group(2)))
            set_run_font(run, size={1: 18, 2: 15, 3: 12.5}[len(heading.group(1))], bold=True, color="111827")
            index += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            set_paragraph_spacing(paragraph, after=3, line_spacing=1.22)
            add_inline_markdown_runs(paragraph, bullet.group(1), size=10.5)
            index += 1
            continue

        numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            set_paragraph_spacing(paragraph, after=3, line_spacing=1.22)
            add_inline_markdown_runs(paragraph, numbered.group(1), size=10.5)
            index += 1
            continue

        quote = re.match(r"^>\s+(.+)$", stripped)
        if quote:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = None
            set_paragraph_spacing(paragraph, before=3, after=6, line_spacing=1.22)
            run = paragraph.add_run("“")
            set_run_font(run, size=11, bold=True, color="5B7CFF")
            add_inline_markdown_runs(paragraph, quote.group(1), size=10.5)
            run = paragraph.add_run("”")
            set_run_font(run, size=11, bold=True, color="5B7CFF")
            index += 1
            continue

        paragraph = document.add_paragraph()
        set_paragraph_spacing(paragraph, after=6, line_spacing=1.28)
        add_inline_markdown_runs(paragraph, stripped, size=10.5)
        index += 1

    if in_code_block:
        flush_code()


def _is_upload_file(value: Any) -> bool:
    """Detect uploaded files from multipart forms.

    FastAPI's request.form() yields starlette UploadFile objects. Checking
    against fastapi.UploadFile with isinstance can miss them and drop files.
    """
    if isinstance(value, UploadFile):
        return True
    return bool(getattr(value, "filename", None) is not None and hasattr(value, "read"))


def collect_upload_files(form: Any) -> list[UploadFile]:
    files: list[UploadFile] = []
    for key, value in form.multi_items():
        if key not in {"files", "file", "attachment", "attachments"}:
            continue
        if not _is_upload_file(value):
            continue
        filename = str(getattr(value, "filename", None) or "").strip()
        # Browsers sometimes include an empty file field with no name.
        if not filename:
            continue
        files.append(value)
    return files


async def parse_chat_job_request(request: Request) -> tuple[str, int | None, str, str | None, list[UploadFile]]:
    content_type = (request.headers.get("content-type") or "").lower()
    if "multipart/form-data" in content_type:
        form = await request.form()
        message = str(form.get("message") or "").strip()
        raw_session_id = str(form.get("session_id") or "").strip()
        provider = normalize_provider(str(form.get("provider") or "openai"))
        model = str(form.get("model") or "").strip() or None
        files = collect_upload_files(form)
        session_id = int(raw_session_id) if raw_session_id.isdigit() else None
        logger.info("Parsed multipart chat job: message_len=%s files=%s", len(message), len(files))
        return message, session_id, provider, model, files

    payload = ChatRequest.model_validate(await request.json())
    return payload.message.strip(), payload.session_id, normalize_provider(payload.provider), payload.model, []


async def save_attachments(
    files: list[UploadFile],
    db: Session,
    user_id: int,
    session_id: int,
    message_id: int,
) -> list[str]:
    if not files:
        return []
    if len(files) > MAX_FILES:
        raise HTTPException(status_code=400, detail=f"Upload at most {MAX_FILES} files at once.")

    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    names: list[str] = []
    saved_paths: list[Path] = []
    try:
        for upload in files:
            original_name = str(upload.filename or "attachment")
            filename = safe_filename(original_name)
            ext = Path(filename).suffix.lower() or Path(original_name).suffix.lower()
            if ext not in ALLOWED_EXTENSIONS:
                raise HTTPException(status_code=400, detail=f"Unsupported attachment type: {original_name}")

            data = await upload.read()
            if not data:
                # Retry once in case the stream was partially consumed.
                try:
                    await upload.seek(0)
                    data = await upload.read()
                except Exception:
                    data = b""
            if not data:
                raise HTTPException(status_code=400, detail=f"Empty attachment: {original_name}")
            if len(data) > MAX_FILE_SIZE:
                raise HTTPException(status_code=400, detail=f"Attachment exceeds 20MB: {original_name}")

            content_type = (getattr(upload, "content_type", None) or "").strip() or "application/octet-stream"
            # Some browsers send empty/octet-stream for images; infer from extension.
            if not content_type.startswith("image/"):
                guessed = IMAGE_MIME_BY_EXT.get(ext)
                if guessed:
                    content_type = guessed
            text_content = extract_text(filename, content_type, data)
            if ext in REQUIRE_EXTRACTED_TEXT_EXTENSIONS and not (text_content or "").strip():
                raise HTTPException(
                    status_code=422,
                    detail=(
                        f"无法从 {original_name} 提取可读正文。"
                        "请确认文件未加密、未损坏，并且不是仅包含扫描图片的文档。"
                    ),
                )

            stored_name = f"{user_id}_{message_id}_{uuid.uuid4().hex}_{filename}"
            stored_path = UPLOAD_DIR / stored_name
            stored_path.write_bytes(data)
            saved_paths.append(stored_path)
            db.add(
                ChatAttachment(
                    user_id=user_id,
                    session_id=session_id,
                    message_id=message_id,
                    filename=filename if filename != "attachment" or not original_name else safe_filename(original_name) or "attachment",
                    content_type=content_type,
                    file_path=str(stored_path.resolve()),
                    file_size=len(data),
                    text_content=text_content,
                )
            )
            names.append(original_name)
            logger.info(
                "Saved attachment message_id=%s name=%s size=%s type=%s text_chars=%s",
                message_id,
                original_name,
                len(data),
                content_type,
                len(text_content or ""),
            )
    except Exception:
        for stored_path in saved_paths:
            try:
                stored_path.unlink(missing_ok=True)
            except OSError:
                logger.warning("Failed to clean up attachment after upload error: %s", stored_path)
        raise
    return names


@router.get("/chat/sessions", response_model=list[ChatSessionOut])
def chat_sessions(db: Session = Depends(get_db), user: UserAccount = Depends(current_user)) -> list[ChatSessionOut]:
    return db.query(ChatSession).filter(ChatSession.user_id == user.id).order_by(desc(ChatSession.updated_at)).limit(10).all()


@router.get("/chat/models", response_model=list[ChatModelOut])
def chat_models(db: Session = Depends(get_db), _user: UserAccount = Depends(current_user)) -> list[ChatModelOut]:
    return list_active_chat_models(db)


@router.post("/chat/export-word")
def export_chat_word(
    payload: ChatExportRequest,
    _user: UserAccount = Depends(current_user),
) -> StreamingResponse:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="python-docx is not installed.") from exc

    content = strip_assistant_markup(payload.content)
    if not content:
        raise HTTPException(status_code=422, detail="No answer content to export.")

    document = Document()
    configure_word_document(document)

    title = document.add_heading("", level=1)
    title_run = title.add_run("AIWeb 回答内容")
    set_run_font(title_run, size=20, bold=True, color="111827")
    set_paragraph_spacing(title, before=0, after=4, line_spacing=1.12)
    meta = document.add_paragraph()
    set_paragraph_spacing(meta, after=12, line_spacing=1.1)
    meta_run = meta.add_run(f"导出时间：{now_utc().strftime('%Y-%m-%d %H:%M:%S UTC')}")
    set_run_font(meta_run, size=9, color="6B7280")
    add_markdown_to_document(document, content)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    filename = f"aiweb-answer-{now_utc().strftime('%Y%m%d-%H%M%S')}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/chat/sessions/{session_id}", response_model=ChatSessionDetail)
def chat_session(
    session_id: int,
    db: Session = Depends(get_db),
    user: UserAccount = Depends(current_user),
) -> ChatSessionDetail:
    session = db.get(ChatSession, session_id)
    if session is None or session.user_id != user.id:
        raise HTTPException(status_code=404, detail="Chat session not found.")
    messages = db.query(ChatMessage).filter(ChatMessage.session_id == session_id).order_by(ChatMessage.created_at).all()
    return ChatSessionDetail(session=session, messages=messages)


@router.delete("/chat/sessions/{session_id}")
def delete_chat_session(
    session_id: int,
    db: Session = Depends(get_db),
    user: UserAccount = Depends(current_user),
) -> dict[str, str]:
    session = db.get(ChatSession, session_id)
    if session is None or session.user_id != user.id:
        raise HTTPException(status_code=404, detail="Chat session not found.")

    attachments = db.query(ChatAttachment).filter(ChatAttachment.session_id == session_id, ChatAttachment.user_id == user.id).all()
    for attachment in attachments:
        try:
            Path(attachment.file_path).unlink(missing_ok=True)
        except OSError:
            pass

    db.query(ChatAttachment).filter(ChatAttachment.session_id == session_id, ChatAttachment.user_id == user.id).delete(synchronize_session=False)
    db.query(ChatJob).filter(ChatJob.session_id == session_id, ChatJob.user_id == user.id).delete(synchronize_session=False)
    db.query(ChatMessage).filter(ChatMessage.session_id == session_id).delete(synchronize_session=False)
    db.delete(session)
    db.commit()
    return {"status": "ok"}


@router.get("/chat/jobs/{job_id}", response_model=ChatJobOut)
def chat_job(job_id: int, db: Session = Depends(get_db), user: UserAccount = Depends(current_user)) -> ChatJobOut:
    job = db.get(ChatJob, job_id)
    if job is None or job.user_id != user.id:
        raise HTTPException(status_code=404, detail="Chat job not found.")
    return job


@router.post("/chat/jobs", response_model=ChatJobOut, dependencies=[Depends(rate_limiter)])
async def create_chat_job(
    request: Request,
    db: Session = Depends(get_db),
    user: UserAccount = Depends(current_user),
) -> ChatJobOut:
    """Enqueue a chat job for the in-process worker. Does not call the model inline."""
    user_message, session_id, provider, requested_model, files = await parse_chat_job_request(request)
    if not user_message and not files:
        raise HTTPException(status_code=422, detail="Please enter a message or upload an attachment.")
    if not user_message and files:
        user_message = "请分析这些附件。"
    if len(user_message) > 4000:
        raise HTTPException(status_code=422, detail="Input cannot exceed 4000 characters.")
    try:
        selected_model = resolve_chat_model(db, provider, requested_model)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc

    session = db.get(ChatSession, session_id) if session_id else None
    if session is not None and session.user_id != user.id:
        raise HTTPException(status_code=404, detail="Chat session not found.")
    if session is None:
        session = ChatSession(title=make_session_title(user_message), user_id=user.id)
        db.add(session)
        db.flush()

    message = ChatMessage(session_id=session.id, role="user", content=user_message)
    db.add(message)
    session.updated_at = now_utc()
    db.flush()

    attachment_names = await save_attachments(files, db, user.id, session.id, message.id)
    if files and not attachment_names:
        raise HTTPException(status_code=400, detail="附件上传失败，请重试。")
    if attachment_names:
        message.content = f"{user_message}\n\nAttachments: {', '.join(attachment_names)}"
        # Ensure the attachment rows and updated message are flushed before commit.
        db.flush()

    job = ChatJob(
        user_id=user.id,
        session_id=session.id,
        user_message_id=message.id,
        provider=provider,
        model=selected_model,
        status="pending",
    )
    db.add(job)
    db.commit()
    db.refresh(job)
    return job


@router.post("/chat", response_model=ChatResponse, dependencies=[Depends(rate_limiter)])
def chat(
    payload: ChatRequest,
    db: Session = Depends(get_db),
    user: UserAccount = Depends(current_user),
) -> ChatResponse:
    """Synchronous chat (compatibility). Prefer POST /api/chat/jobs for production UI."""
    user_message = payload.message.strip()
    session = db.get(ChatSession, payload.session_id) if payload.session_id else None
    if session is not None and session.user_id != user.id:
        raise HTTPException(status_code=404, detail="Chat session not found.")
    if session is None:
        session = ChatSession(title=make_session_title(user_message), user_id=user.id)
        db.add(session)
        db.flush()

    history = load_recent_chat_history(db, session.id)
    provider = normalize_provider(payload.provider)
    try:
        selected_model = resolve_chat_model(db, provider, payload.model)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc

    try:
        service = OpenAIService(provider=provider, text_model=selected_model)
        result = service.chat(user_message, history=history)
        text = str(result["text"])
        db.add(ChatMessage(session_id=session.id, role="user", content=user_message))
        db.add(ChatMessage(session_id=session.id, role="assistant", content=text))
        db.add(ChatRecord(user_id=user.id, user_message=user_message, ai_response=text))
        record_token_usage(
            db,
            user_id=user.id,
            source="chat",
            provider=provider,
            model=str(result.get("model") or service.text_model),
            prompt_tokens=int(result.get("prompt_tokens") or 0),
            completion_tokens=int(result.get("completion_tokens") or 0),
            total_tokens=int(result.get("total_tokens") or 0),
        )
        session.updated_at = now_utc()
        db.commit()
        return ChatResponse(text=text, session_id=session.id)
    except OpenAIServiceError as exc:
        db.rollback()
        raise HTTPException(status_code=502, detail=str(exc)) from exc
